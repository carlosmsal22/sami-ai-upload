# /pages/Executive_Insight_Generator.py
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import base64
import tempfile

st.set_page_config(page_title="Executive Insight Generator", layout="wide")
st.title("📊 Executive Insight Generator")

st.sidebar.header("Upload Crosstab File")
uploaded_file = st.sidebar.file_uploader("Upload a WinCross-style Excel file (.xlsx)", type="xlsx")

if uploaded_file:
    sheet_names = pd.ExcelFile(uploaded_file).sheet_names
    selected_sheet = st.sidebar.selectbox("Select Sheet", sheet_names)
    df = pd.read_excel(uploaded_file, sheet_name=selected_sheet, header=None)
    st.success(f"Loaded sheet: {selected_sheet}")

    # Custom dropdown from banner plan
    banner_options = [
        "Total",
        "Residential Units Managed",
        "Job Role",
        "Job Title",
        "Change in Living Expectations",
        "Track Utility Expenses",
        "Percent of Rent Delinquent",
        "Security Deposit Return Method",
        "Staff Turnover Rate",
        "Digital vs Paper Rent Payments"
    ]
    selected_banner = st.sidebar.selectbox("Select Banner Breakout", banner_options)

    # --- Helper Functions ---
    def parse_tables(df):
        segment_names = ["Frugal Basics", "Resourceful Savers", "Performance Enthusiasts", "Urban Techies"]
        tables = []
        row = 1
        while row < len(df):
            if "Table Title" in str(df.iloc[row, 1]):
                table_title = str(df.iloc[row + 1, 1])
                base_counts = df.iloc[row + 6, 3:7].tolist()
                segment_labels = [
                    f"{name} (n={int(count)})" if pd.notnull(count) else f"{name} (n=NA)"
                    for name, count in zip(segment_names, base_counts)
                ]
                table_rows = []
                sub_row = row + 8
                while sub_row + 2 < len(df) and isinstance(df.iloc[sub_row, 2], str):
                    metric_label = df.iloc[sub_row, 2]
                    try:
                        freqs = [df.iloc[sub_row, col] for col in range(3, 7)]
                        percs = [df.iloc[sub_row + 1, col] for col in range(3, 7)]
                        sigs_raw = df.iloc[sub_row + 2, 3:7].tolist()
                        values = [
                            f"{float(p)*100:.1f}% ({int(f)})"
                            if pd.notna(p) and pd.notna(f) and p != '-' and f != '-'
                            else ""
                            for p, f in zip(percs, freqs)
                        ]
                        sig_combined = ', '.join([str(sig) for sig in sigs_raw if pd.notna(sig) and isinstance(sig, str)])
                        table_rows.append([metric_label] + values + [sig_combined])
                    except:
                        break
                    sub_row += 3
                table_df = pd.DataFrame(table_rows, columns=["Metric"] + segment_labels + ["Sig"])
                tables.append((table_title, table_df))
                row = sub_row
            else:
                row += 1
        return tables

    def generate_insights(df_table):
        summary_lines = []
        rows = df_table.head(4)
        for _, row in rows.iterrows():
            try:
                values = row[1:-1].tolist()
                metric = row["Metric"]
                high_idx = max(range(len(values)), key=lambda i: float(values[i].split('%')[0]) if '%' in values[i] else -1)
                low_idx = min(range(len(values)), key=lambda i: float(values[i].split('%')[0]) if '%' in values[i] else float('inf'))
                segments = df_table.columns[1:-1].tolist()
                high_val = values[high_idx]
                low_val = values[low_idx]
                summary_lines.append(f"{segments[high_idx]} leads in {metric.lower()} at {high_val}, while {segments[low_idx]} trails at {low_val}.")
            except:
                continue
        return summary_lines

    def generate_chart(df_table, title):
        chart_df = df_table.head(4).copy()
        for col in chart_df.columns[1:-1]:
            chart_df[col] = chart_df[col].str.extract(r"([\d\.]+)%").astype(float)
        fig, ax = plt.subplots(figsize=(8, 4))
        for col in chart_df.columns[1:-1]:
            ax.plot(chart_df["Metric"], chart_df[col], marker='o', label=col)
        ax.set_title(title)
        ax.set_ylabel("%")
        ax.legend()
        plt.xticks(rotation=45, ha="right")
        plt.tight_layout()
        return fig

    tables = parse_tables(df)
    table_titles = [f"Table {i+1}: {t[0][:60]}" for i, t in enumerate(tables)]
    selected_idx = st.sidebar.selectbox("Select a table to preview", options=range(len(tables)), format_func=lambda x: table_titles[x])

    # Display selected table
    table_title, table_df = tables[selected_idx]
    st.subheader(f"📘 {table_title}")
    st.dataframe(table_df, use_container_width=True)

    # Show insights
    st.markdown("**Key Findings:**")
    insights = generate_insights(table_df)
    for line in insights:
        st.markdown(f"- {line}")

    # Show chart
    st.markdown("**Chart:**")
    fig = generate_chart(table_df, table_title)
    st.pyplot(fig)

    # Export options
    if st.sidebar.button("📄 Download Word Report"):
        doc = Document()
        doc.add_heading("SAMI AI Executive Insights Report", level=1)
        for title, df in tables:
            doc.add_heading(title, level=2)
            insight_lines = generate_insights(df)
            doc.add_paragraph("Key Findings:", style="List Bullet")
            for line in insight_lines:
                doc.add_paragraph(line, style="List Bullet 2")
            doc.add_paragraph("Implications:", style="List Bullet")
            doc.add_paragraph("Further strategic implications can be drawn from the segment behaviors and differences above.", style="List Bullet 2")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            with open(tmp.name, "rb") as file:
                st.sidebar.download_button("Download Word File", file.read(), file_name="Executive_Insights_Report.docx")

    if st.sidebar.button("📄 Download PDF Report"):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        for title, df in tables:
            pdf.set_font("Arial", 'B', 12)
            pdf.multi_cell(0, 10, title)
            insight_lines = generate_insights(df)
            pdf.set_font("Arial", size=11)
            for line in insight_lines:
                pdf.multi_cell(0, 10, f"- {line}")
            pdf.multi_cell(0, 10, "Implications:")
            pdf.multi_cell(0, 10, "Further strategic implications can be drawn from the segment behaviors and differences above.")
            pdf.ln(4)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            pdf.output(tmp.name)
            with open(tmp.name, "rb") as file:
                st.sidebar.download_button("Download PDF File", file.read(), file_name="Executive_Insights_Report.pdf")